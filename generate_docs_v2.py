
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_detailed_technical_document():
    document = Document()

    # --- helper to set cell background color ---
    def set_cell_bg(cell, color_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    # --- Styles ---
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Clean H1
    h1 = document.styles['Heading 1']
    h1.font.name = 'Calibri Light'
    h1.font.size = Pt(24)
    h1.font.color.rgb = RGBColor(41, 98, 255) # NeuroStock Blue
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Clean H2
    h2 = document.styles['Heading 2']
    h2.font.name = 'Calibri Light'
    h2.font.size = Pt(18)
    h2.font.color.rgb = RGBColor(0, 150, 136) # Teal
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    
    # Clean H3
    h3 = document.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(66, 66, 66) # Dark Gray

    # --- Title Page ---
    document.add_heading('NeuroStock', 0)
    document.add_heading('NSE Stock Trend Predictor AI', 1)
    
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('\n\nComprehensive Technical Design Specification').bold = True
    p.add_run(f'\n\nVersion: 2.0')
    p.add_run(f'\nDate: {datetime.now().strftime("%B %d, %Y")}')
    p.add_run('\n\n\n\n\n\n\n\n')
    p.add_run('Confidential - Intellectual Property of Alyster Benedict')
    
    document.add_page_break()

    # --- Table of Contents Placeholder (Word auto-updates usually, but we list sections) ---
    document.add_heading('Table of Contents', level=1)
    p = document.add_paragraph()
    p.add_run("1. Executive Summary\n")
    p.add_run("2. System Architecture & Tech Stack\n")
    p.add_run("3. Machine Learning Ecosystem\n")
    p.add_run("    3.1 Data Pipeline & Feature Engineering\n")
    p.add_run("    3.2 Model Architectures (LSTM, XGBoost, etc.)\n")
    p.add_run("    3.3 Train/Test Methodology\n")
    p.add_run("4. Feature Modules (The 9 Core Tabs)\n")
    p.add_run("5. AI Integration (LLM Layer)\n")
    p.add_run("6. Technical Appendix (API Specs)")
    
    document.add_page_break()

    # --- 1. Executive Summary ---
    document.add_heading('1. Executive Summary', level=1)
    document.add_paragraph(
        "NeuroStock uses a hybrid approach to financial forecasting by combining quantitative Machine Learning (ML) models with qualitative Generative AI insights. "
        "Unlike traditional predictors that rely solely on linear regression, NeuroStock deploys a suite of Deep Learning (LSTM) and Ensemble (XGBoost, Random Forest) models "
        "trained on over 10 years of historical OHCLV (Open, High, Low, Close, Volume) data from the National Stock Exchange (NSE)."
    )
    document.add_paragraph(
        "Key Differentiators:\n"
        "- Dual-Forecast System: Separate models for high-accuracy 'Next Day' prediction vs. speculative 'Long Range' trend mapping.\n"
        "- Feature Augmentation: Integration of technical indicators like RSI (Relative Strength Index) and SMA (Simple Moving Average) directly into the feature vectors.\n"
        "- Local AI Privacy: Utilization of LM Studio to run Large Language Models locally, ensuring financial queries remain private."
    )

    # --- 2. System Architecture ---
    document.add_heading('2. System Architecture', level=1)
    document.add_paragraph(
        "The system follows a decoupled Client-Server architecture with a RESTful API layer."
    )

    # Table for Stack
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology Choice'
    set_cell_bg(hdr_cells[0], 'E0E0E0')
    set_cell_bg(hdr_cells[1], 'E0E0E0')

    specs = [
        ('Frontend Framework', 'React 18 (Functional Components, Hooks)'),
        ('Charting Engine', 'Lightweight Charts (Canvas-based, by TradingView)'),
        ('UI Styling', 'CSS Modules, Glassmorphism Design System'),
        ('Backend API', 'Flask (Python 3.9+)'),
        ('ML Library', 'TensorFlow 2.x (Keras), Scikit-Learn, XGBoost'),
        ('Data Processing', 'Pandas, NumPy, Joblib (Serialization)'),
        ('Data Source', 'yfinance (Yahoo Finance API)'),
        ('Auth Provider', 'Google Firebase Authentication'),
        ('LLM Server', 'LM Studio (Local Inference Server)')
    ]

    for comp, tech in specs:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = tech

    document.add_heading('2.1 Data Flow Diagram', level=2)
    document.add_paragraph(
        "1. Client (React) initiates a request (e.g., 'Predict Infosys for Dec 25').\n"
        "2. Flask Server validates the token and request payload.\n"
        "3. Loader Module checks for existing pre-trained models (`.h5` or `.joblib`) on disk.\n"
        "4. Data Fetcher retrieves live OHLCV data via `yfinance` to build the context window (last 100 days).\n"
        "5. ML Inference Engine scales the data, runs the prediction, and inverse-scales the result.\n"
        "6. (Optional) AI Module sends the trend data to LM Studio to generate a textual narrative.\n"
        "7. Response is sent back to Client as JSON."
    )

    # --- 3. ML Ecosystem ---
    document.add_heading('3. Machine Learning Ecosystem', level=1)
    document.add_paragraph("The core intelligence of NeuroStock facilitates two distinct prediction modes:")
    
    document.add_heading('3.1 Mode A: High-Fidelity 1-Day Forecast', level=2)
    document.add_paragraph(
        "Designed for maximum accuracy for the immediate next trading day. It uses a 'Feature-Rich' approach."
    )
    document.add_paragraph(
        "Input Features (Vector Size: 7):"
    )
    p = document.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Open, High, Low, Close, Volume\n")
    p.add_run("SMA_50 (50-day Simple Moving Average)\n")
    p.add_run("RSI_14 (14-day Relative Strength Index)")

    document.add_heading('3.2 Mode B: Long-Range Speculative Forecast', level=2)
    document.add_paragraph(
        "Uses a 'Walk-Forward' algorithm to predict prices deep into the future. Limitations in error propagation mean this is strictly estimation."
    )
    document.add_paragraph(
        "Input Features (Vector Size: 1): Close Price only. (To reduce noise accumulation over recursive steps)."
    )

    document.add_heading('3.3 Model Architectures', level=2)
    
    document.add_heading('3.3.1 LSTM (Deep Learning)', level=3)
    document.add_paragraph(
        "The Long Short-Term Memory network is the primary engine. Architecture:"
    )
    document.add_paragraph(
        "- Input Layer: Shape (Time_Step=100, Features=7 or 1)\n"
        "- Hidden Layer 1: LSTM (50 Units, Return Sequences=True)\n"
        "- Dropout: 0.2 (20% drop to prevent overfitting)\n"
        "- Hidden Layer 2: LSTM (50 Units, Return Sequences=False)\n"
        "- Dropout: 0.2\n"
        "- Dense Layer: 25 Units (Relu activation implicit)\n"
        "- Output Layer: 1 Unit (Linear activation for regression)"
    )

    document.add_heading('3.3.2 Ensemble Regressors', level=3)
    document.add_paragraph(
        "To provide consensus, the system also trains 5 traditional ML models on flattened time-series data:"
    )
    document.add_paragraph(
        "1. XGBoost: Gradient boosting decision tree (n_estimators=100).\n"
        "2. Random Forest: Bagging ensemble (n_estimators=100).\n"
        "3. SVR: Support Vector Regressor (Kernel='rbf', C=1.0).\n"
        "4. Decision Tree: Single tree baseline.\n"
        "5. Linear Regression: Baseline to detect pure linear trends."
    )

    # --- 4. The 9 Tabs ---
    document.add_heading('4. Feature Modules (The 9 Core Tabs)', level=1)
    
    # Tab 1
    document.add_heading('4.1 Market Dashboard (Home)', level=2)
    document.add_paragraph(
        "The command center for market awareness."
    )
    document.add_paragraph(
        "Technical Implementation:\n"
        "- Endpoint: `/get-market-data`\n"
        "- Logic: Fetches live quotes for all Nifty 50 tickers in parallel.\n"
        "- UI Optimization: Uses a `TickerTape` component with CSS keyframe animations for smooth scrolling without JS overhead.\n"
        "- Search: Client-side fuzzy search to filter the stock grid instantly."
    )

    # Tab 2
    document.add_heading('4.2 AI Investment Engine', level=2)
    document.add_paragraph(
        "Automated Portfolio Construction tool."
    )
    document.add_paragraph(
        "Technical Implementation:\n"
        "- Endpoint: `/investment-engine` (POST)\n"
        "- Simulation: The backend iterates through the top 30 stocks. For each stock, it runs a prediction simulation up to the user's `Withdrawal Date`.\n"
        "- Ranking Logic: Stocks are ranked by `(Predicted_Appreciation / Volatility_Index)`.\n"
        "- Output: Returns the top 3 candidates with an AI-generated 'Investment Logic' explanation."
    )

    # Tab 3
    document.add_heading('4.3 Prediction (The "Neuro" Core)', level=2)
    document.add_paragraph(
        "Single-Asset deep dive visualization."
    )
    document.add_paragraph(
        "Technical Implementation:\n"
        "- Endpoint: `/predict` (POST)\n"
        "- Dynamic Graphing: The frontend receives two arrays: `Historical_Prices` (Static) and `Forecast_Prices` (Dynamic). These are spliced together and rendered on a Canvas element using `Lightweight-Charts`.\n"
        "- Color Logic: The forecast line color is determined conditionally: `Green` if Final > Current, else `Red`."
    )

    # Tab 4
    document.add_heading('4.4 Model Comparison', level=2)
    document.add_paragraph(
        "A/B Validation Suite."
    )
    document.add_paragraph(
        "Technical Implementation:\n"
        "- Concurrency: The frontend issues a unified request containing `[Algo_A, Algo_B]`. The backend uses Python's threading or sequential processing to run inference on both loaded models.\n"
        "- Visual Diff: Two distinct chart cards are rendered side-by-side, allowing pixel-perfect comparison of the trend curve."
    )

    # Tab 5
    document.add_heading('4.5 Portfolio Projector', level=2)
    document.add_paragraph(
        "Personalized Wealth Calculator."
    )
    document.add_paragraph(
        "Technical Implementation:\n"
        "- Logic: `Future_Value = Current_Value * (Predicted_Price_T / Current_Price_0)`\n"
        "- State: Persistent local state allows users to switch between stocks without losing their entered principal amounts."
    )

    # Tab 6
    document.add_heading('4.6 Positions Dashboard', level=2)
    document.add_paragraph(
        "Leaderboard of Market Movers."
    )
    document.add_paragraph(
        "Technical Implementation:\n"
        "- Sorting Algorithms: Client-side QuickSort implementation to sort the table by 'Day Change %', 'Price', or 'Volume'.\n"
        "- Visuals: Conditional CSS classes (`.pos`, `.neg`) apply standardized color tokens (Emerald Green, Crimson Red) across the application."
    )

    # Tab 7
    document.add_heading('4.7 52-Week Analysis', level=2)
    document.add_paragraph(
        "Long-term Support/Resistance visualization."
    )
    document.add_paragraph(
        "Technical Implementation:\n"
        "- Data Range: Fetches `1y` data specifically via `yfinance`.\n"
        "- Overlays: Calculates and renders static horizontal lines for `High_Max` and `Low_Min` on the chart surface."
    )

    # Tab 8
    document.add_heading('4.8 Sentiment Dashboard', level=2)
    document.add_paragraph(
        "NLP-driven Qualitative Analysis."
    )
    document.add_paragraph(
        "Technical Implementation:\n"
        "- Endpoint: `/get-sentiment`\n"
        "- Data Source: Scrapes/fetches recent news headlines for the ticker.\n"
        "- Scoring: Words are tokenized and scored against a financial lexicon (or via LLM prompt) to derive a `Sentiment_Score` (0-100).\n"
        "- Visualization: A custom SVG Gauge component renders the score dynamically."
    )

    # Tab 9
    document.add_heading('4.9 Financial Teacher (FinTeach)', level=2)
    document.add_paragraph(
        "Educational Chat Interface."
    )
    document.add_paragraph(
        "Technical Implementation:\n"
        "- Context Separation: Unlike the other AI tools, FinTeach has a system prompt enforcing a 'Professor' persona. It explicitly denies answering questions about live stock prices to avoid hallucination.\n"
        "- Interaction: Standard Chat UX with 'Typing...' indicators and scroll-to-bottom behavior."
    )

    # --- 5. AI Integration ---
    document.add_heading('5. AI Integration (Local LLM)', level=1)
    document.add_paragraph(
        "The application interfaces with a locally running Large Language Model to ensure data privacy."
    )
    document.add_paragraph(
        "Prompt Engineering Strategy:\n"
        "The system constructs dynamic prompts injecting quantitative data into the context window. Example structure:"
    )
    
    # Use a table for the code block to allow background shading
    code_table = document.add_table(rows=1, cols=1)
    code_cell = code_table.rows[0].cells[0]
    set_cell_bg(code_cell, 'F5F5F5') # Light gray background
    
    code_ex = code_cell.paragraphs[0]
    code_ex.style = 'No Spacing'
    run = code_ex.add_run(
        'System: You are a financial analyst.\n'
        'User: Analyze Infosys. RSI is 65. SMA is Trending Up. Volatility is Low.\n'
        'Assistant: [Generates Insight...]'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # --- 6. Appendix ---
    document.add_heading('6. Appendix: API Specification', level=1)
    
    table2 = document.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Description'
    
    apis = [
        ('POST', '/predict', 'Main forecasting engine. Returns trend JSON.'),
        ('POST', '/investment-engine', 'Runs multi-stock simulation.'),
        ('GET', '/get-market-data', 'Returns live quotes for all tickers.'),
        ('POST', '/get-ai-insights', 'Invokes LLM for specific stock analysis.'),
        ('POST', '/get-general-knowledge', 'Invokes LLM for educational Q&A.')
    ]
    
    for method, endp, desc in apis:
        row = table2.add_row().cells
        row[0].text = method
        row[1].text = endp
        row[2].text = desc

    # Save
    file_path = 'NeuroStock_Comprehensive_Technical_Spec.docx'
    document.save(file_path)
    print(f"Document saved to {os.path.abspath(file_path)}")

if __name__ == "__main__":
    try:
        create_detailed_technical_document()
    except Exception as e:
        print(f"Error: {e}")
